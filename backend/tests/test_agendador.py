"""O agendador e a exportação — o bloco que faz o sistema trabalhar sozinho.

O que se testa aqui não é "a função roda". É o que quebra um agendador de
verdade: rodar duas vezes no mesmo dia e duplicar tudo, disparar o mesmo alerta
a cada varredura, uma fonte fora do ar derrubar a rodada inteira.
"""
from __future__ import annotations

import datetime as dt

import pytest
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.agendador.tarefas import executar, rodar_prazos, rodar_vigilias
from app.core.db import Base
from app.core.enums import Prioridade, StatusPrazo
from app.engines.exportacao.motor import (
    MARCADOR_LACUNA,
    classificar,
    nome_de_arquivo,
    para_docx,
    para_pdf,
)
from app.models import Base as ModelosBase  # noqa: F401
from app.modules.agendador.models import ExecucaoTarefa
from app.modules.normativo.models import MonitoramentoNormativo
from app.modules.prazos.models import Notificacao, Pendencia, Prazo
from app.seeds.carga import popular


@pytest.fixture
def db(tmp_path):
    engine = create_engine(f"sqlite+pysqlite:///{tmp_path / 'agenda.db'}")
    Base.metadata.create_all(engine)
    Sessao = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    with Sessao() as sessao:
        popular(sessao, com_demonstracao=True)
        yield sessao


# ------------------------------------------------------------------ prazos


def test_varredura_de_prazos_materializa_a_agenda(db):
    relatorio = rodar_prazos(db)
    prazos = db.scalars(select(Prazo)).all()

    assert prazos, "a entidade de demonstração tem mandato e certidões com data"
    assert relatorio.numeros.get("prazos abertos") == len(prazos)
    # Prazo nunca é presumido: todo um tem origem e a data que o gerou (§21).
    assert all(p.origem in ("LEI", "ESTATUTO", "RCPJ", "MANUAL") for p in prazos)
    assert all(p.chave_idempotencia for p in prazos)


def test_rodar_duas_vezes_no_mesmo_dia_nao_duplica_nada(db):
    rodar_prazos(db)
    quantos = len(db.scalars(select(Prazo)).all())
    notificacoes = len(db.scalars(select(Notificacao)).all())
    pendencias = len(db.scalars(select(Pendencia)).all())

    segunda = rodar_prazos(db)

    assert len(db.scalars(select(Prazo)).all()) == quantos
    assert len(db.scalars(select(Notificacao)).all()) == notificacoes, (
        "o mesmo alerta não pode ser reenviado a cada varredura"
    )
    assert len(db.scalars(select(Pendencia)).all()) == pendencias
    assert "prazos abertos" not in segunda.numeros


def test_alerta_dispara_uma_vez_por_janela_e_de_novo_na_janela_seguinte(db):
    """Cada faixa (90, 60, 30...) avisa uma vez. Aproximar-se avisa de novo."""
    rodar_prazos(db)
    prazo = db.scalars(
        select(Prazo).where(Prazo.status == StatusPrazo.ABERTO).order_by(Prazo.data_limite)
    ).first()
    assert prazo is not None

    db.query(Notificacao).delete()
    prazo.alertas_disparados = []
    prazo.janelas_alerta = [90, 30, 7]
    db.add(prazo)
    db.commit()

    trinta_dias_antes = prazo.data_limite - dt.timedelta(days=30)
    rodar_prazos(db, hoje=trinta_dias_antes)
    apos_primeira = len(db.scalars(select(Notificacao)).all())
    assert apos_primeira > 0

    rodar_prazos(db, hoje=trinta_dias_antes + dt.timedelta(days=1))
    assert len(db.scalars(select(Notificacao)).all()) == apos_primeira, (
        "ainda na janela dos 30 dias: não avisa de novo"
    )

    rodar_prazos(db, hoje=prazo.data_limite - dt.timedelta(days=7))
    assert len(db.scalars(select(Notificacao)).all()) > apos_primeira, (
        "entrou na janela dos 7 dias: é um alerta novo, mais urgente"
    )


def test_prazo_que_sai_da_agenda_para_de_alertar(db):
    rodar_prazos(db)
    prazo = db.scalars(select(Prazo)).first()
    prazo.chave_idempotencia = "CHAVE_QUE_A_AGENDA_NAO_GERA_MAIS"
    db.add(prazo)
    db.commit()

    rodar_prazos(db)
    db.refresh(prazo)
    assert prazo.status is StatusPrazo.CUMPRIDO


def test_prazo_remarcado_volta_a_alertar(db):
    """Mandato prorrogado: os alertas da data antiga não valem para a nova."""
    rodar_prazos(db)
    prazo = db.scalars(select(Prazo).order_by(Prazo.data_limite)).first()
    prazo.data_limite = prazo.data_limite - dt.timedelta(days=400)
    prazo.alertas_disparados = [90, 60, 30]
    db.add(prazo)
    db.commit()

    rodar_prazos(db)
    db.refresh(prazo)
    assert prazo.alertas_disparados != [90, 60, 30]


# ---------------------------------------------------------------- vigílias


def test_vigilia_so_verifica_fonte_vencida(db):
    from app.core.tempo import agora

    monitoramentos = db.scalars(select(MonitoramentoNormativo)).all()
    assert monitoramentos, "a carga inicial cadastra as vigílias da base legal"
    for m in monitoramentos:
        m.ultima_verificacao = agora()
        db.add(m)
    db.commit()

    relatorio = rodar_vigilias(db)
    assert relatorio.numeros == {}, "todas em dia: nada a verificar"


def test_fonte_manual_vencida_cobra_o_responsavel_sem_repetir(db):
    from app.core.tempo import agora

    for m in db.scalars(select(MonitoramentoNormativo)).all():
        m.modo = "MANUAL"
        m.ultima_verificacao = agora() - dt.timedelta(days=400)
        db.add(m)
    db.commit()

    primeira = rodar_vigilias(db)
    assert primeira.numeros["cobranças de conferência manual"] > 0
    abertas = len(db.scalars(select(Pendencia)).all())

    rodar_vigilias(db)
    assert len(db.scalars(select(Pendencia)).all()) == abertas, (
        "a mesma cobrança em aberto não vira uma pendência nova a cada dia"
    )


def test_fonte_que_falha_nao_derruba_a_rodada(db, monkeypatch):
    from app.core.tempo import agora
    import app.agendador.tarefas as tarefas

    monitoramentos = db.scalars(select(MonitoramentoNormativo)).all()
    for m in monitoramentos:
        m.modo = "HTTP"
        m.ultima_verificacao = agora() - dt.timedelta(days=400)
        db.add(m)
    db.commit()

    chamadas = {"n": 0}

    def falhar_na_primeira(db_, monitoramento, coleta=None):
        chamadas["n"] += 1
        if chamadas["n"] == 1:
            raise RuntimeError("site do Planalto fora do ar")
        return None

    monkeypatch.setattr(tarefas, "verificar_monitoramento", falhar_na_primeira)
    relatorio = rodar_vigilias(db)

    assert len(relatorio.falhas) == 1
    assert relatorio.resultado == "PARCIAL"
    assert relatorio.numeros["fontes verificadas"] == len(monitoramentos) - 1


# --------------------------------------------------------------- execução


def test_execucao_fica_registrada_para_auditoria(db):
    registro = executar(db, "prazos", acionada_por="MANUAL")
    assert registro.resultado == "OK"
    assert registro.concluida_em is not None
    assert registro.duracao_segundos >= 0
    assert registro.acionada_por == "MANUAL"
    assert db.scalars(select(ExecucaoTarefa)).all()


def test_falha_geral_tambem_fica_registrada(db, monkeypatch):
    import app.agendador.tarefas as tarefas

    def explodir(*_args, **_kw):
        raise RuntimeError("banco caiu no meio da varredura")

    monkeypatch.setitem(tarefas.TAREFAS, "PRAZOS", explodir)
    registro = executar(db, "prazos")

    assert registro.resultado == "ERRO"
    assert "banco caiu" in registro.detalhe
    assert registro.falhas


def test_tarefa_desconhecida_e_recusada(db):
    with pytest.raises(ValueError):
        executar(db, "faxina")


# -------------------------------------------------------------- exportação


TEXTO = """ASSOCIAÇÃO COMUNITÁRIA NOVO HORIZONTE
CNPJ nº 12.345.678/0001-90

ATA DA ASSEMBLEIA GERAL ORDINÁRIA

Aos doze dias do mês de março de dois mil e vinte e seis, reuniram-se os
associados em pleno gozo de seus direitos estatutários.

ORDEM DO DIA

1. Prestação de contas;
2. Eleição da diretoria;

O Tesoureiro, Sr. """ + MARCADOR_LACUNA + """, apresentou as contas.

Belo Horizonte, 12 de março de 2026.


_______________________________________________
Maria Aparecida Souza
Presidente
"""


def test_estrutura_do_documento_e_reconhecida():
    especies = [b.especie for b in classificar(TEXTO)]
    assert especies[0] == "CABECALHO"
    assert "TITULO" in especies
    assert especies.count("ITEM") == 2, "cada item da ordem do dia é um parágrafo"
    assert "FECHO" in especies, "local e data vão alinhados à direita"
    assert especies[-1] == "NOME_ASSINATURA"


def test_os_dois_formatos_leem_a_mesma_estrutura():
    docx = para_docx(TEXTO, "Ata")
    pdf = para_pdf(TEXTO, "Ata")
    assert docx[:2] == b"PK", "docx é um zip"
    assert pdf[:5] == b"%PDF-"


def test_lacuna_sai_destacada_no_docx(tmp_path):
    from docx import Document

    caminho = tmp_path / "ata.docx"
    caminho.write_bytes(para_docx(TEXTO, "Ata"))
    documento = Document(str(caminho))

    marcadas = [
        run.text
        for p in documento.paragraphs
        for run in p.runs
        if run.bold and run.font.color and run.font.color.rgb is not None
    ]
    assert "DADO NÃO INFORMADO" in marcadas, (
        "documento incompleto que sai bonitinho é o que alguém protocola sem ver"
    )
    assert MARCADOR_LACUNA not in "\n".join(p.text for p in documento.paragraphs)


def test_docx_sai_no_formato_forense(tmp_path):
    from docx import Document

    caminho = tmp_path / "ata.docx"
    caminho.write_bytes(para_docx(TEXTO, "Ata"))
    documento = Document(str(caminho))
    secao = documento.sections[0]

    assert round(secao.page_width.cm, 1) == 21.0
    assert round(secao.page_height.cm, 1) == 29.7
    assert round(secao.left_margin.cm) == 3
    assert documento.styles["Normal"].font.size.pt == 12


def test_nome_do_arquivo_atravessa_email_e_cartorio():
    assert nome_de_arquivo("Ata: Eleição de diretoria — 2026", "docx") == (
        "Ata-Eleicao-de-diretoria-2026.docx"
    )
    assert nome_de_arquivo("///", "pdf") == "documento.pdf"

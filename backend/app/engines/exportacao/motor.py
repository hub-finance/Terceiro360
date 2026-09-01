"""EXPORTAÇÃO DE DOCUMENTOS — DOCX e PDF (§19, §28).

O documento gerado é texto corrido com as convenções de uma peça jurídica:
cabeçalho com a qualificação da entidade, título em caixa alta, corpo
justificado, e ao final as linhas de assinatura. Exportar não é despejar essas
linhas num arquivo — é reconhecer essa estrutura e dar a ela a forma que o
cartório espera receber.

Duas decisões de projeto:

1. **Uma estrutura, dois desenhistas.** `classificar()` decide uma vez o que é
   cabeçalho, título, item, fecho e assinatura; DOCX e PDF apenas desenham esse
   mesmo resultado. O risco de o PDF sair diferente do arquivo editável não vem
   de ter dois formatos — vem de ter duas leituras do documento. Aqui há uma só.
   (Converter o DOCX com LibreOffice daria fidelidade maior ainda, mas ao preço
   de exigir a suíte inteira instalada no servidor para exportar uma ata.)

2. **Lacuna exportada continua gritando.** O marcador DADO NÃO INFORMADO sai em
   negrito e vermelho nos dois formatos — cor que sobrevive à impressão em preto
   e branco como cinza escuro destacado. Um documento incompleto que sai
   bonitinho é pior do que um que não sai: alguém protocola sem perceber (§46).
"""
from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

MARCADOR_LACUNA = "**DADO NÃO INFORMADO**"

# ABNT/uso forense: A4, margens 3-2-3-2, Times 12, entrelinha 1,5.
FONTE = "Times New Roman"
CORPO_PT = 12
MARGENS_CM = {"superior": 3, "inferior": 2, "esquerda": 3, "direita": 2}

_REGUA_ASSINATURA = re.compile(r"^[_\-–—]{5,}$")
_ITEM_NUMERADO = re.compile(r"^\s*(\d+[.)]|[a-z][.)]|[IVXLC]+\s*[-–.)])\s+")
# "São Paulo, 12 de março de 2026." — o fecho de local e data, que por convenção
# forense vai alinhado à direita, logo acima das assinaturas.
_FECHO_LOCAL_DATA = re.compile(
    r"^.{2,60},\s*\d{1,2}\s+de\s+[^\d]{3,20}\s+de\s+\d{4}\.?$", re.IGNORECASE
)


VERMELHO = (0xB3, 0x20, 0x2B)

# Fontes serifadas com cobertura Unicode, em ordem de preferência. A Liberation
# Serif tem as mesmas métricas da Times New Roman declarada no DOCX, então as
# duas exportações quebram linha no mesmo lugar.
FONTES_PDF = (
    ("Liberation Serif", "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
     "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf"),
    ("DejaVu Serif", "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"),
    ("FreeSerif", "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
     "/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf"),
)

# A Times embutida no PDF só conhece latin-1. Sem fonte Unicode instalada, um
# travessão derrubaria a exportação inteira — trocar o sinal por seu equivalente
# ASCII estraga a tipografia, e é muito melhor do que não entregar o arquivo.
SEM_LATIN1 = {
    "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u00a0": " ",
    "\u2022": "-", "\u2212": "-", "\u2032": "'", "\u2033": '"',
}


def _fonte_unicode() -> tuple[str, str, str] | None:
    for nome, regular, negrito in FONTES_PDF:
        if Path(regular).exists() and Path(negrito).exists():
            return nome, regular, negrito
    return None


def _reduzir_a_latin1(texto: str) -> str:
    for original, substituto in SEM_LATIN1.items():
        texto = texto.replace(original, substituto)
    # O que ainda não couber vira "?" em vez de derrubar a exportação.
    return texto.encode("latin-1", "replace").decode("latin-1")


@dataclass
class Bloco:
    """Um parágrafo do documento, já classificado."""

    texto: str
    especie: str  # CABECALHO|TITULO|ITEM|FECHO|CORPO|REGUA|NOME_ASSINATURA


def _e_caixa_alta(linha: str) -> bool:
    letras = [c for c in linha if c.isalpha()]
    if not letras:
        return False
    return all(c.isupper() for c in letras)


def classificar(texto: str) -> list[Bloco]:
    """Descobre a estrutura do documento a partir das convenções do texto.

    Nenhuma heurística aqui inventa conteúdo: ela só decide alinhamento e peso
    da fonte. Errar a classificação deixa o documento feio, nunca errado.
    """
    linhas = [linha.rstrip() for linha in texto.replace("\r\n", "\n").split("\n")]

    # Jinja deixa linhas em branco onde havia `{% if %}` e `{% for %}`. Três
    # linhas vazias seguidas viram uma só: o buraco é do gerador, não do autor.
    limpas: list[str] = []
    for linha in linhas:
        if not linha.strip() and limpas and not limpas[-1].strip():
            continue
        limpas.append(linha)

    blocos: list[Bloco] = []
    paragrafo: list[str] = []
    # O cabeçalho vai até a primeira linha em branco: é o bloco de qualificação
    # da entidade, que sempre abre a peça.
    no_cabecalho = True
    assinatura_a_seguir = False

    def fechar():
        nonlocal paragrafo
        if not paragrafo:
            return
        junto = " ".join(p.strip() for p in paragrafo).strip()
        paragrafo = []
        if not junto:
            return
        if no_cabecalho:
            especie = "CABECALHO"
        elif assinatura_a_seguir:
            especie = "NOME_ASSINATURA"
        elif _ITEM_NUMERADO.match(junto):
            especie = "ITEM"
        elif _e_caixa_alta(junto) and len(junto) <= 120:
            especie = "TITULO"
        elif _FECHO_LOCAL_DATA.match(junto):
            especie = "FECHO"
        else:
            especie = "CORPO"
        blocos.append(Bloco(junto, especie))

    for linha in limpas:
        crua = linha.strip()
        if not crua:
            fechar()
            if no_cabecalho and blocos:
                no_cabecalho = False
            continue
        if _REGUA_ASSINATURA.match(crua):
            fechar()
            blocos.append(Bloco("", "REGUA"))
            # Quem vem depois da régua é o nome de quem assina, não corpo.
            assinatura_a_seguir = True
            continue
        # Cada item numerado é um bloco: emendar "1. …" com "2. …" num
        # parágrafo só transformaria a ordem do dia numa frase corrida.
        if _ITEM_NUMERADO.match(crua):
            fechar()
            paragrafo = [crua]
            fechar()
            continue
        # Cabeçalho e títulos são de uma linha; corpo se junta em parágrafo.
        if no_cabecalho or assinatura_a_seguir or _e_caixa_alta(crua):
            fechar()
            paragrafo = [crua]
            fechar()
            continue
        paragrafo.append(crua)

    fechar()
    return blocos


def _aplicar_texto(paragrafo, texto: str) -> None:
    """Escreve o texto marcando as lacunas em negrito e realce amarelo."""
    for pedaco in re.split(f"({re.escape(MARCADOR_LACUNA)})", texto):
        if not pedaco:
            continue
        run = paragrafo.add_run(
            "DADO NÃO INFORMADO" if pedaco == MARCADOR_LACUNA else pedaco
        )
        if pedaco == MARCADOR_LACUNA:
            run.bold = True
            run.font.color.rgb = RGBColor(0xB3, 0x20, 0x2B)
            # Sem realce nativo em cor de fundo aqui: o vermelho em negrito
            # sobrevive a impressão em preto e branco, o amarelo não.


def para_docx(texto: str, titulo: str, rodape: str | None = None) -> bytes:
    """Monta o .docx e devolve os bytes."""
    documento = Document()

    secao = documento.sections[0]
    secao.orientation = WD_ORIENT.PORTRAIT
    secao.page_width, secao.page_height = Cm(21), Cm(29.7)
    secao.top_margin = Cm(MARGENS_CM["superior"])
    secao.bottom_margin = Cm(MARGENS_CM["inferior"])
    secao.left_margin = Cm(MARGENS_CM["esquerda"])
    secao.right_margin = Cm(MARGENS_CM["direita"])

    normal = documento.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(CORPO_PT)

    documento.core_properties.title = titulo
    documento.core_properties.comments = (
        "Gerado pelo TERCEIRO360. Minuta sujeita a revisão de profissional habilitado."
    )

    for bloco in classificar(texto):
        if bloco.especie == "REGUA":
            regua = documento.add_paragraph()
            regua.alignment = WD_ALIGN_PARAGRAPH.CENTER
            regua.paragraph_format.space_before = Pt(24)
            regua.paragraph_format.space_after = Pt(0)
            regua.add_run("_" * 48)
            continue

        p = documento.add_paragraph()
        formato = p.paragraph_format
        formato.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if bloco.especie == "CABECALHO":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
            formato.space_after = Pt(0)
        elif bloco.especie == "TITULO":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.space_before = Pt(18)
            formato.space_after = Pt(6)
        elif bloco.especie == "NOME_ASSINATURA":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
            formato.space_after = Pt(0)
        elif bloco.especie == "FECHO":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            formato.space_before = Pt(18)
        elif bloco.especie == "ITEM":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            formato.left_indent = Cm(1.25)
            formato.space_after = Pt(4)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            formato.first_line_indent = Cm(1.25)
            formato.space_after = Pt(6)

        _aplicar_texto(p, bloco.texto)
        if bloco.especie == "TITULO":
            for run in p.runs:
                run.bold = True

    if rodape:
        p = documento.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(rodape)
        run.font.size = Pt(8)
        run.font.name = FONTE

    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def para_pdf(texto: str, titulo: str, rodape: str | None = None) -> bytes:
    """Monta o .pdf a partir da mesma estrutura que gera o .docx."""
    from fpdf import FPDF

    escolhida = _fonte_unicode()
    familia = escolhida[0] if escolhida else "Times"

    def texto_seguro(conteudo: str) -> str:
        return conteudo if escolhida else _reduzir_a_latin1(conteudo)

    class Peca(FPDF):
        def footer(self):
            if not rodape:
                return
            self.set_y(-15)
            self.set_font(familia, size=8)
            self.set_text_color(0x66, 0x66, 0x66)
            self.cell(0, 5, texto_seguro(f"{rodape} · página {self.page_no()}"), align="C")

    pdf = Peca(format="A4", unit="mm")
    if escolhida:
        pdf.add_font(familia, "", escolhida[1])
        pdf.add_font(familia, "B", escolhida[2])
    pdf.set_margins(
        left=MARGENS_CM["esquerda"] * 10,
        top=MARGENS_CM["superior"] * 10,
        right=MARGENS_CM["direita"] * 10,
    )
    pdf.set_auto_page_break(True, margin=MARGENS_CM["inferior"] * 10)
    pdf.set_title(titulo)
    pdf.add_page()
    pdf.set_font(familia, size=CORPO_PT)

    def escrever(paragrafo, conteudo: str, negrito: bool = False) -> None:
        """Escreve marcando as lacunas — a cor é o alarme (§46)."""
        for pedaco in re.split(f"({re.escape(MARCADOR_LACUNA)})", conteudo):
            if not pedaco:
                continue
            lacuna = pedaco == MARCADOR_LACUNA
            pdf.set_font(familia, style="B" if (negrito or lacuna) else "", size=CORPO_PT)
            pdf.set_text_color(*(VERMELHO if lacuna else (0, 0, 0)))
            paragrafo.write(texto_seguro("DADO NÃO INFORMADO" if lacuna else pedaco))
        pdf.set_font(familia, size=CORPO_PT)
        pdf.set_text_color(0, 0, 0)

    # mm — o docx trabalha em cm, o fpdf em mm; o recuo é o mesmo 1,25 cm.
    RECUO = 12.5
    ESTILOS = {
        "CABECALHO": dict(text_align="C", line_height=1.15, bottom_margin=0),
        "TITULO": dict(text_align="C", line_height=1.15, top_margin=6, bottom_margin=2),
        "ITEM": dict(text_align="J", line_height=1.5, indent=RECUO, bottom_margin=1.5),
        "FECHO": dict(text_align="R", line_height=1.5, top_margin=6),
        "NOME_ASSINATURA": dict(text_align="C", line_height=1.15, bottom_margin=0),
        "CORPO": dict(text_align="J", line_height=1.5, first_line_indent=RECUO,
                      bottom_margin=2),
    }

    with pdf.text_columns() as coluna:
        for bloco in classificar(texto):
            if bloco.especie == "REGUA":
                with coluna.paragraph(text_align="C", top_margin=8, bottom_margin=0) as p:
                    p.write("_" * 48)
                continue
            estilo = ESTILOS.get(bloco.especie, ESTILOS["CORPO"])
            with coluna.paragraph(**estilo) as p:
                escrever(p, bloco.texto, negrito=bloco.especie == "TITULO")

    return bytes(pdf.output())


def nome_de_arquivo(titulo: str, extensao: str) -> str:
    """Nome previsível e seguro para anexar em e-mail ou levar ao cartório.

    Sem acento de propósito: o nome viaja num cabeçalho HTTP, por e-mail e por
    pendrive até o balcão do cartório, e cada etapa dessas tem seu próprio jeito
    de estragar um "ç". O conteúdo do documento é acentuado; o nome do arquivo
    não precisa ser.
    """
    sem_acento = unicodedata.normalize("NFKD", titulo).encode("ascii", "ignore").decode()
    limpo = re.sub(r"[^A-Za-z0-9\s-]", "", sem_acento).strip()
    limpo = re.sub(r"[\s_]+", "-", limpo).strip("-")
    return f"{(limpo or 'documento')[:80]}.{extensao}"
